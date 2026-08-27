from __future__ import annotations

import json
import logging
import threading
from types import SimpleNamespace

import httpx
import pytest
from fastapi import HTTPException


def _response(*, arguments: str = '{"topic":"light"}') -> SimpleNamespace:
    return SimpleNamespace(
        output=[
            SimpleNamespace(
                type="message",
                content=[SimpleNamespace(type="output_text", text="Draft ready")],
            ),
            SimpleNamespace(
                type="function_call",
                call_id="call-1",
                name="emit_document",
                arguments=arguments,
            ),
        ],
        usage=SimpleNamespace(
            input_tokens=17,
            output_tokens=5,
            input_tokens_details=SimpleNamespace(cached_tokens=3),
        ),
    )


class _FakeFiles:
    def __init__(self, events: list[tuple], *, delete_failures: set[str] | None = None):
        self.events = events
        self.delete_failures = delete_failures or set()
        self.created = 0

    def create(self, **request):
        self.created += 1
        if self.created == _FakeArk.create_failure_at:
            raise RuntimeError("upload failed")
        remote_id = f"remote-{self.created}"
        self.events.append(("create", remote_id, request))
        return SimpleNamespace(id=remote_id)

    def delete(self, remote_id: str, **kwargs):
        assert kwargs == {"timeout": 2.0}
        self.events.append(("delete", remote_id))
        if remote_id in self.delete_failures:
            raise RuntimeError(f"provider body contains {remote_id}")


class _FakeResponses:
    def __init__(self, events: list[tuple], response, *, failure: BaseException | None = None):
        self.events = events
        self.response = response
        self.failure = failure

    def create(self, **request):
        self.events.append(("response", request))
        if self.failure is not None:
            raise self.failure
        return self.response


class _FakeArk:
    instances: list[_FakeArk] = []
    response = _response()
    response_failure: BaseException | None = None
    delete_failures: set[str] = set()
    create_failure_at = 0

    def __init__(self, **kwargs):
        self.kwargs = kwargs
        self.events: list[tuple] = []
        self.files = _FakeFiles(self.events, delete_failures=self.delete_failures)
        self.responses = _FakeResponses(
            self.events,
            self.response,
            failure=self.response_failure,
        )
        self.instances.append(self)


@pytest.fixture(autouse=True)
def _reset_fake_ark(monkeypatch):
    from app.gateway import provider_file_cleanup

    _FakeArk.instances = []
    _FakeArk.response = _response()
    _FakeArk.response_failure = None
    _FakeArk.delete_failures = set()
    _FakeArk.create_failure_at = 0
    monkeypatch.setattr(
        provider_file_cleanup,
        "enqueue_provider_file_cleanup",
        lambda provider, connection, remote: f"cleanup-{remote}",
    )
    monkeypatch.setattr(
        provider_file_cleanup,
        "mark_provider_file_cleanup_succeeded",
        lambda cleanup_id: None,
    )


def _assets():
    from app.gateway.ark_responses import ResolvedDocumentAsset

    return [
        ResolvedDocumentAsset(
            resource_id="doc-1",
            filename="normal.pdf",
            content_type="application/pdf",
            content=b"%PDF-normal",
        ),
        ResolvedDocumentAsset(
            resource_id="doc-2",
            filename="scan.pdf",
            content_type="application/pdf",
            content=b"%PDF-scanned-image-pages",
        ),
    ]


def _payload():
    return {
        "systemPrompt": "Keep the planning boundary.",
        "messages": [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": "Compare these references"},
                    {"type": "image", "mimeType": "image/png", "data": "aW1hZ2U="},
                ],
            },
            {
                "role": "assistant",
                "content": [
                    {"type": "text", "text": "I will inspect them."},
                    {
                        "type": "toolCall",
                        "id": "prior-call",
                        "name": "lookup",
                        "arguments": {"key": "value"},
                    },
                ],
            },
            {"role": "toolResult", "toolCallId": "prior-call", "content": "done"},
        ],
        "tools": [
            {
                "name": "emit_document",
                "description": "Emit a draft",
                "parameters": {"type": "object", "properties": {}},
            }
        ],
        "temperature": 0.2,
        "maxTokens": 123,
    }


def test_files_responses_and_delete_are_ordered_and_mixed_input_is_normalized(monkeypatch):
    from app.gateway import ark_responses

    monkeypatch.setattr(ark_responses, "Ark", _FakeArk)

    result = ark_responses.complete_ark_response(
        _payload(),
        api_base="https://ark.example/api/v3",
        credential="secret",
        model_id="doubao-seed-2-0-lite-260215",
        pdf_assets=_assets(),
        connection_id="connection-1",
    )

    instance = _FakeArk.instances[0]
    assert [event[0] for event in instance.events] == [
        "create", "create", "response", "delete", "delete"
    ]
    request = instance.events[2][1]
    assert request["instructions"] == "Keep the planning boundary."
    assert request["model"] == "doubao-seed-2-0-lite-260215"
    assert request["store"] is False
    assert request["max_output_tokens"] == 123
    assert request["tools"] == [{
        "type": "function",
        "name": "emit_document",
        "description": "Emit a draft",
        "parameters": {"type": "object", "properties": {}},
    }]
    assert request["input"][0] == {
        "role": "user",
        "content": [
            {"type": "input_text", "text": "Compare these references"},
            {
                "type": "input_image",
                "detail": "auto",
                "image_url": "data:image/png;base64,aW1hZ2U=",
            },
            {"type": "input_file", "file_id": "remote-1"},
            {"type": "input_file", "file_id": "remote-2"},
        ],
    }
    assert request["input"][1:] == [
        {"role": "assistant", "content": "I will inspect them."},
        {
            "type": "function_call",
            "call_id": "prior-call",
            "name": "lookup",
            "arguments": '{"key": "value"}',
        },
        {"type": "function_call_output", "call_id": "prior-call", "output": "done"},
    ]
    assert result == {
        "content": [
            {"type": "text", "text": "Draft ready"},
            {
                "type": "toolCall",
                "id": "call-1",
                "name": "emit_document",
                "arguments": {"topic": "light"},
            },
        ],
        "stopReason": "toolUse",
        "usage": {"input": 17, "output": 5, "cacheRead": 3, "cacheWrite": 0},
    }


@pytest.mark.parametrize("failure", [RuntimeError("model failed"), KeyboardInterrupt()])
def test_every_uploaded_file_is_deleted_when_response_aborts(monkeypatch, failure):
    from app.gateway import ark_responses

    _FakeArk.response_failure = failure
    monkeypatch.setattr(ark_responses, "Ark", _FakeArk)

    with pytest.raises(type(failure)):
        ark_responses.complete_ark_response(
            _payload(),
            api_base="https://ark.example/api/v3",
            credential="secret",
            model_id="model-1",
            pdf_assets=_assets(),
            connection_id="connection-1",
        )

    assert [event[:2] for event in _FakeArk.instances[0].events if event[0] == "delete"] == [
        ("delete", "remote-1"),
        ("delete", "remote-2"),
    ]


def test_provider_model_error_is_rethrown_as_a_redacted_code(monkeypatch, caplog):
    from app.gateway import ark_responses

    _FakeArk.response_failure = RuntimeError(
        "raw-body remote-1 C:/private/document.pdf secret-credential"
    )
    monkeypatch.setattr(ark_responses, "Ark", _FakeArk)

    with pytest.raises(Exception) as error:
        ark_responses.complete_ark_response(
            _payload(),
            api_base="https://ark.example/api/v3",
            credential="secret-credential",
            model_id="model-1",
            pdf_assets=_assets(),
            connection_id="connection-1",
        )

    assert str(error.value) == "provider_response_failed"
    assert "raw-body" not in caplog.text
    assert "remote-1" not in caplog.text
    assert "private/document.pdf" not in caplog.text
    assert "secret-credential" not in caplog.text


def test_delete_failure_leaves_pre_registered_cleanup_intents_pending(monkeypatch, caplog):
    from app.gateway import ark_responses

    _FakeArk.delete_failures = {"remote-1", "remote-2"}
    monkeypatch.setattr(ark_responses, "Ark", _FakeArk)
    enqueued: list[tuple[str, str, str]] = []

    def register(provider: str, connection: str, remote: str) -> str:
        enqueued.append((provider, connection, remote))
        return f"cleanup-{remote}"

    result = ark_responses.complete_ark_response(
        _payload(),
        api_base="https://ark.example/api/v3",
        credential="secret",
        model_id="model-1",
        pdf_assets=_assets(),
        connection_id="connection-1",
        enqueue_cleanup=register,
    )

    assert result["content"][0]["text"] == "Draft ready"
    assert enqueued == [
        ("volcengine-ark", "connection-1", "remote-1"),
        ("volcengine-ark", "connection-1", "remote-2"),
    ]
    logged = caplog.text
    assert "remote-1" not in logged
    assert "remote-2" not in logged
    assert "secret" not in logged
    assert "%PDF" not in logged


def test_partial_upload_failure_deletes_every_successfully_uploaded_file(monkeypatch):
    from app.gateway import ark_responses

    _FakeArk.create_failure_at = 2
    monkeypatch.setattr(ark_responses, "Ark", _FakeArk)

    with pytest.raises(RuntimeError, match="provider_file_upload_failed"):
        ark_responses.complete_ark_response(
            _payload(),
            api_base="https://ark.example/api/v3",
            credential="secret",
            model_id="model-1",
            pdf_assets=_assets(),
            connection_id="connection-1",
        )

    assert [event[:2] for event in _FakeArk.instances[0].events] == [
        ("create", "remote-1"),
        ("delete", "remote-1"),
    ]


def test_malformed_response_tool_arguments_normalize_to_empty_object(monkeypatch):
    from app.gateway import ark_responses

    _FakeArk.response = _response(arguments="not-json")
    monkeypatch.setattr(ark_responses, "Ark", _FakeArk)

    result = ark_responses.complete_ark_response(
        _payload(),
        api_base="https://ark.example/api/v3",
        credential="secret",
        model_id="model-1",
        pdf_assets=_assets()[:1],
        connection_id="connection-1",
    )

    assert result["content"][1]["arguments"] == {}


class _CleanupStorage:
    def __init__(self, due: list[dict]):
        self.due = due
        self.succeeded: list[str] = []
        self.retries: list[tuple[str, int, str]] = []

    def get_due(self, *, now: int, limit: int):
        assert now == 1_000_000
        assert limit <= 20
        return list(self.due)

    def mark_succeeded(self, cleanup_id: str):
        self.succeeded.append(cleanup_id)

    def mark_retry(self, cleanup_id: str, *, next_attempt_at: int, error_code: str):
        self.retries.append((cleanup_id, next_attempt_at, error_code))


class _CleanupCredentials:
    def __init__(self, values: dict[str, str]):
        self.values = values
        self.lookups: list[str] = []

    def get(self, connection_id: str):
        self.lookups.append(connection_id)
        return self.values.get(connection_id)


class _CleanupConnections:
    def __init__(self, credentials: _CleanupCredentials):
        self.credential_store = credentials

    def get_connection_config(self, connection_id: str):
        return {
            "id": connection_id,
            "providerId": "volcengine-ark",
            "apiBase": "https://ark.example/api/v3",
            "enabled": True,
        }


class _CleanupArk:
    deleted: list[tuple[str, str, str]] = []
    failure: BaseException | None = None

    def __init__(
        self,
        *,
        api_key: str,
        base_url: str,
        timeout: float,
        max_retries: int,
    ):
        self.api_key = api_key
        self.base_url = base_url
        assert 0 < timeout <= 2.0
        assert max_retries == 0
        self.files = self

    def delete(self, remote_file_id: str):
        self.deleted.append((self.api_key, self.base_url, remote_file_id))
        if self.failure:
            raise self.failure


def _due_cleanup():
    return {
        "cleanupId": "cleanup-1",
        "providerId": "volcengine-ark",
        "connectionId": "connection-1",
        "remoteFileId": "opaque-remote-id",
        "createdAt": 10,
        "lastAttemptAt": None,
        "attemptCount": 0,
        "nextAttemptAt": 10,
        "lastErrorCode": None,
    }


def test_cleanup_retry_uses_fresh_credential_and_marks_success(monkeypatch, caplog):
    from app.gateway.provider_file_cleanup import retry_provider_file_cleanup

    storage = _CleanupStorage([_due_cleanup()])
    credentials = _CleanupCredentials({"connection-1": "fresh-secret"})
    connections = _CleanupConnections(credentials)
    _CleanupArk.deleted = []
    _CleanupArk.failure = None

    summary = retry_provider_file_cleanup(
        limit=20,
        storage=storage,
        connection_store=connections,
        ark_factory=_CleanupArk,
        now_ms=lambda: 1_000_000,
    )

    assert summary.attempted == 1
    assert summary.succeeded == 1
    assert summary.retry_scheduled == 0
    assert credentials.lookups == ["connection-1"]
    assert _CleanupArk.deleted == [
        ("fresh-secret", "https://ark.example/api/v3", "opaque-remote-id")
    ]
    assert storage.succeeded == ["cleanup-1"]
    assert storage.retries == []
    assert "opaque-remote-id" not in caplog.text
    assert "fresh-secret" not in caplog.text


def test_cleanup_retry_persists_redacted_code_and_is_bounded(monkeypatch, caplog):
    from app.gateway.provider_file_cleanup import retry_provider_file_cleanup

    storage = _CleanupStorage([_due_cleanup(), {**_due_cleanup(), "cleanupId": "cleanup-2"}])
    credentials = _CleanupCredentials({"connection-1": "fresh-secret"})
    _CleanupArk.deleted = []
    _CleanupArk.failure = RuntimeError(
        "raw provider body opaque-remote-id C:/private/file.pdf fresh-secret"
    )

    summary = retry_provider_file_cleanup(
        limit=1,
        storage=storage,
        connection_store=_CleanupConnections(credentials),
        ark_factory=_CleanupArk,
        now_ms=lambda: 1_000_000,
    )

    assert summary.attempted == 1
    assert summary.succeeded == 0
    assert summary.retry_scheduled == 1
    assert storage.retries == [("cleanup-1", 1_060_000, "provider_cleanup_failed")]
    assert len(_CleanupArk.deleted) == 1
    assert "opaque-remote-id" not in caplog.text
    assert "fresh-secret" not in caplog.text
    assert "private/file.pdf" not in caplog.text


def test_cleanup_intent_is_durable_before_response_and_survives_marker_outage(
    monkeypatch,
):
    from app.gateway import ark_responses

    registered: list[tuple[str, str, str, str]] = []

    def register(provider: str, connection: str, remote: str) -> str:
        cleanup_id = f"cleanup-{len(registered) + 1}"
        registered.append((cleanup_id, provider, connection, remote))
        _FakeArk.instances[0].events.append(("register", remote))
        return cleanup_id

    def unavailable_marker(cleanup_id: str) -> None:
        raise RuntimeError(f"storage unavailable for {cleanup_id}")

    monkeypatch.setattr(ark_responses, "Ark", _FakeArk)
    result = ark_responses.complete_ark_response(
        _payload(),
        api_base="https://ark.example/api/v3",
        credential="secret",
        model_id="model-1",
        pdf_assets=_assets(),
        connection_id="connection-1",
        enqueue_cleanup=register,
        mark_cleanup_succeeded=unavailable_marker,
    )

    assert result["content"][0]["text"] == "Draft ready"
    assert registered == [
        ("cleanup-1", "volcengine-ark", "connection-1", "remote-1"),
        ("cleanup-2", "volcengine-ark", "connection-1", "remote-2"),
    ]
    assert [event[0] for event in _FakeArk.instances[0].events] == [
        "create",
        "register",
        "create",
        "register",
        "response",
        "delete",
        "delete",
    ]


def test_cleanup_registration_failure_aborts_model_and_returns_only_fixed_code(
    monkeypatch,
    caplog,
):
    from app.gateway import ark_responses

    remote_file_id = "remote-sensitive-registration"

    def fail_registration(provider: str, connection: str, remote: str) -> str:
        raise RuntimeError(
            f"storage raw body {remote_file_id} C:/private.pdf secret-credential"
        )

    monkeypatch.setattr(ark_responses, "Ark", _FakeArk)
    with pytest.raises(ark_responses.ArkResponsesError) as error:
        ark_responses.complete_ark_response(
            _payload(),
            api_base="https://ark.example/api/v3",
            credential="secret-credential",
            model_id="model-1",
            pdf_assets=_assets()[:1],
            connection_id="connection-1",
            enqueue_cleanup=fail_registration,
        )

    assert str(error.value) == "provider_cleanup_persistence_failed"
    assert all(event[0] != "response" for event in _FakeArk.instances[0].events)
    assert [event[0] for event in _FakeArk.instances[0].events] == ["create", "delete"]
    assert remote_file_id not in caplog.text
    assert "private.pdf" not in caplog.text
    assert "secret-credential" not in caplog.text


def test_cleanup_retry_treats_already_absent_remote_file_as_idempotent_success():
    from app.gateway.provider_file_cleanup import retry_provider_file_cleanup

    class AlreadyAbsentError(RuntimeError):
        status_code = 404

    storage = _CleanupStorage([_due_cleanup()])
    credentials = _CleanupCredentials({"connection-1": "fresh-secret"})
    _CleanupArk.deleted = []
    _CleanupArk.failure = AlreadyAbsentError("raw provider body with remote id")

    summary = retry_provider_file_cleanup(
        storage=storage,
        connection_store=_CleanupConnections(credentials),
        ark_factory=_CleanupArk,
        now_ms=lambda: 1_000_000,
    )

    assert summary.succeeded == 1
    assert summary.retry_scheduled == 0
    assert storage.succeeded == ["cleanup-1"]


def test_cleanup_retry_bounds_provider_delete_and_stops_at_batch_deadline():
    from app.gateway.provider_file_cleanup import retry_provider_file_cleanup

    storage = _CleanupStorage(
        [_due_cleanup(), {**_due_cleanup(), "cleanupId": "cleanup-2"}]
    )
    credentials = _CleanupCredentials({"connection-1": "fresh-secret"})
    ark_options: list[dict] = []

    class BoundedArk:
        def __init__(self, **kwargs):
            ark_options.append(kwargs)
            self.files = self

        def delete(self, remote_file_id: str):
            assert remote_file_id == "opaque-remote-id"

    clock = iter([10.0, 10.1, 12.1])
    summary = retry_provider_file_cleanup(
        storage=storage,
        connection_store=_CleanupConnections(credentials),
        ark_factory=BoundedArk,
        now_ms=lambda: 1_000_000,
        monotonic=lambda: next(clock),
        batch_budget_seconds=2.0,
    )

    assert summary.attempted == 1
    assert ark_options == [
        {
            "api_key": "fresh-secret",
            "base_url": "https://ark.example/api/v3",
            "timeout": pytest.approx(1.9),
            "max_retries": 0,
        }
    ]


def test_cleanup_storage_client_uses_authenticated_internal_contract(monkeypatch):
    from app.gateway.provider_file_cleanup import ProviderCleanupStorageClient

    requests: list[httpx.Request] = []

    def handler(request: httpx.Request):
        requests.append(request)
        if request.url.path.endswith("/due"):
            return httpx.Response(200, json={"items": [_due_cleanup()]})
        if request.url.path == "/api/internal/provider-file-cleanup":
            return httpx.Response(200, json={"cleanupId": "cleanup-1"})
        return httpx.Response(200, json={"ok": True})

    monkeypatch.setenv("PROMPTCARD_INTERNAL_TOKEN", "internal-token")
    http = httpx.Client(
        base_url="http://storage.test",
        transport=httpx.MockTransport(handler),
    )
    client = ProviderCleanupStorageClient(client=http)

    assert client.get_due(now=1_000_000, limit=20) == [_due_cleanup()]
    assert (
        client.enqueue("volcengine-ark", "connection-1", "opaque-remote-id")
        == "cleanup-1"
    )
    client.mark_succeeded("cleanup-1")
    client.mark_retry(
        "cleanup-1",
        next_attempt_at=1_060_000,
        error_code="provider_cleanup_failed",
    )

    assert [request.url.path for request in requests] == [
        "/api/internal/provider-file-cleanup/due",
        "/api/internal/provider-file-cleanup",
        "/api/internal/provider-file-cleanup/succeeded",
        "/api/internal/provider-file-cleanup/retry",
    ]
    assert all(not request.url.query for request in requests[2:])
    assert requests[2].read() == b'{"cleanupId":"cleanup-1"}'
    assert json.loads(requests[3].read()) == {
        "cleanupId": "cleanup-1",
        "nextAttemptAt": 1_060_000,
        "errorCode": "provider_cleanup_failed",
    }
    assert all(
        request.headers["X-PromptCard-Internal-Token"] == "internal-token"
        for request in requests
    )


def test_cleanup_storage_client_keeps_cleanup_id_out_of_request_target(monkeypatch):
    from app.gateway.provider_file_cleanup import ProviderCleanupStorageClient

    requests: list[httpx.Request] = []
    cleanup_id = "cleanup-sensitive/../private.pdf?token=secret"
    monkeypatch.setenv("PROMPTCARD_INTERNAL_TOKEN", "internal-token")
    client = ProviderCleanupStorageClient(
        client=httpx.Client(
            base_url="http://storage.test",
            transport=httpx.MockTransport(
                lambda request: requests.append(request)
                or httpx.Response(200, json={"ok": True})
            ),
        )
    )

    client.mark_succeeded(cleanup_id)
    client.mark_retry(
        cleanup_id,
        next_attempt_at=1_060_000,
        error_code="provider_cleanup_failed",
    )

    assert [request.url.raw_path for request in requests] == [
        b"/api/internal/provider-file-cleanup/succeeded",
        b"/api/internal/provider-file-cleanup/retry",
    ]
    assert [json.loads(request.read()) for request in requests] == [
        {"cleanupId": cleanup_id},
        {
            "cleanupId": cleanup_id,
            "nextAttemptAt": 1_060_000,
            "errorCode": "provider_cleanup_failed",
        },
    ]


def test_gateway_suppresses_http_client_urls_with_provider_and_cleanup_ids(
    caplog,
):
    from app.gateway.app import create_app
    from app.gateway.provider_file_cleanup import ProviderCleanupStorageClient

    remote_file_id = "remote-file-sensitive-3"
    cleanup_id = "cleanup-sensitive-7"
    transport = httpx.MockTransport(
        lambda request: httpx.Response(200, json={"ok": True})
    )
    raw_client = httpx.Client(
        base_url="https://provider.invalid",
        transport=transport,
    )
    storage = ProviderCleanupStorageClient(client=raw_client)

    create_app()
    with caplog.at_level(logging.INFO):
        raw_client.delete(f"/files/{remote_file_id}")
        storage.mark_succeeded(cleanup_id)

    assert remote_file_id not in caplog.text
    assert cleanup_id not in caplog.text


def _descriptor(*, provider: str = "volcengine-ark", inputs=None):
    return {
        "connectionId": "connection-1",
        "providerId": provider,
        "model": {
            "id": "doubao-seed-2-0-lite-260215",
            "displayName": "Doubao Seed 2.0 Lite",
            "modality": "chat",
            "capabilities": {"input": inputs or ["text", "image", "pdf"]},
            "integrationGroup": {
                "id": "volcengine-ark-sdk",
                "displayName": "Ark SDK",
                "kind": "sdk",
            },
        },
    }


@pytest.mark.anyio
async def test_document_ids_are_resolved_before_model_and_handle_is_request_scoped(monkeypatch):
    from app.gateway import promptcard_runtime
    from app.gateway.ark_responses import ResolvedDocumentAsset

    assets = [
        ResolvedDocumentAsset(
            resource_id="text-1",
            filename="notes.txt",
            content_type="text/plain",
            content=b"notes",
            normalized_text="notes",
        ),
        ResolvedDocumentAsset(
            resource_id="pdf-1",
            filename="scan.pdf",
            content_type="application/pdf",
            content=b"%PDF-scan",
        ),
    ]
    order: list[str] = []
    captured_handle = ""

    async def load_documents(project_id, resource_ids):
        order.append("documents")
        assert project_id == "project-1"
        assert resource_ids == ["text-1", "pdf-1"]
        return assets

    def resolve_model(binding):
        order.append("model")
        return _descriptor()

    async def invoke(payload):
        nonlocal captured_handle
        order.append("invoke")
        assert "documentResourceIds" not in payload
        assert "explicitDocumentNodeIds" not in payload
        captured_handle = payload["documentInvocationHandle"]
        first = promptcard_runtime._document_invocations.resolve(captured_handle)
        second = promptcard_runtime._document_invocations.resolve(captured_handle)
        assert first is second
        assert [asset.resource_id for asset in first.assets] == ["pdf-1"]
        return {"text": "ok", "canvasEdits": [], "proposals": []}

    monkeypatch.setattr(promptcard_runtime, "_load_document_resources", load_documents)
    monkeypatch.setattr(
        promptcard_runtime,
        "require_pdf_text_model",
        lambda binding: _descriptor(),
    )
    monkeypatch.setattr(promptcard_runtime, "resolve_text_model", resolve_model)
    monkeypatch.setattr(promptcard_runtime, "_invoke_text_agent", invoke)
    monkeypatch.setattr(promptcard_runtime, "_load_canvas_image_attachments", _empty_attachments)
    body = promptcard_runtime.PromptCardRuntimeMessageRequest.model_validate({
        "content": "Read these",
        "projectId": "project-1",
        "permissionScope": "prompt-library-agent",
        "documentResourceIds": ["text-1", "pdf-1"],
    })

    result = await promptcard_runtime.runtime_service.send_message(body, None)

    assert result["text"] == "ok"
    assert order == ["documents", "model", "invoke"]
    assert promptcard_runtime._document_invocations.resolve(captured_handle) is None


@pytest.mark.anyio
async def test_text_agent_disconnect_discards_request_scoped_pdf_handle(monkeypatch):
    from app.gateway import promptcard_runtime
    from app.gateway.ark_responses import ResolvedDocumentAsset

    async def load_documents(project_id, resource_ids):
        return [ResolvedDocumentAsset(
            resource_id="pdf-1",
            filename="scan.pdf",
            content_type="application/pdf",
            content=b"%PDF-scan",
        )]

    captured_handle = ""

    async def disconnect(payload):
        nonlocal captured_handle
        captured_handle = payload["documentInvocationHandle"]
        assert promptcard_runtime._document_invocations.resolve(captured_handle) is not None
        raise HTTPException(status_code=503, detail="text_agent_unavailable")

    monkeypatch.setattr(promptcard_runtime, "_load_document_resources", load_documents)
    monkeypatch.setattr(
        promptcard_runtime,
        "require_pdf_text_model",
        lambda binding: _descriptor(),
    )
    monkeypatch.setattr(promptcard_runtime, "resolve_text_model", lambda binding: _descriptor())
    monkeypatch.setattr(promptcard_runtime, "_invoke_text_agent", disconnect)
    monkeypatch.setattr(promptcard_runtime, "_load_canvas_image_attachments", _empty_attachments)
    body = promptcard_runtime.PromptCardRuntimeMessageRequest.model_validate({
        "content": "Read",
        "projectId": "project-1",
        "permissionScope": "prompt-library-agent",
        "documentResourceIds": ["pdf-1"],
    })

    with pytest.raises(HTTPException, match="text_agent_unavailable"):
        await promptcard_runtime.runtime_service.send_message(body, None)

    assert promptcard_runtime._document_invocations.resolve(captured_handle) is None


async def _empty_attachments(*args, **kwargs):
    return []


@pytest.mark.anyio
@pytest.mark.parametrize(
    ("resource_ids", "detail"),
    [
        (["same", "same"], "document_resource_ids_duplicate"),
        (["one", "two"], "document_attachments_too_large"),
    ],
)
async def test_invalid_document_sets_reject_before_model_resolution(
    monkeypatch,
    resource_ids,
    detail,
):
    from app.gateway import promptcard_runtime
    from app.gateway.ark_responses import ResolvedDocumentAsset

    monkeypatch.setattr(promptcard_runtime, "MAX_DOCUMENT_TURN_BYTES", 10)
    calls: list[str] = []

    async def load_documents(project_id, ids):
        calls.append("documents")
        return [
            ResolvedDocumentAsset(
                resource_id=value,
                filename=f"{value}.txt",
                content_type="text/plain",
                content=b"123456",
                normalized_text="123456",
            )
            for value in ids
        ]

    def resolve_model(binding):
        calls.append("model")
        return _descriptor()

    monkeypatch.setattr(promptcard_runtime, "_load_document_resources", load_documents)
    monkeypatch.setattr(promptcard_runtime, "resolve_text_model", resolve_model)
    body = promptcard_runtime.PromptCardRuntimeMessageRequest.model_validate({
        "content": "Read",
        "projectId": "project-1",
        "permissionScope": "prompt-library-agent",
        "documentResourceIds": resource_ids,
    })

    with pytest.raises(HTTPException) as error:
        await promptcard_runtime.runtime_service.send_message(body, None)

    assert error.value.detail == detail
    assert "model" not in calls


@pytest.mark.anyio
async def test_pdf_capability_rejects_before_text_agent_invocation(monkeypatch):
    from app.gateway import promptcard_runtime
    from app.gateway.ark_responses import ResolvedDocumentAsset

    async def load_documents(project_id, ids):
        return [ResolvedDocumentAsset(
            resource_id="pdf-1",
            filename="scan.pdf",
            content_type="application/pdf",
            content=b"%PDF-scan",
        )]

    invoked = False

    async def invoke(payload):
        nonlocal invoked
        invoked = True

    monkeypatch.setattr(promptcard_runtime, "_load_document_resources", load_documents)
    monkeypatch.setattr(
        promptcard_runtime,
        "require_pdf_text_model",
        lambda binding: (_ for _ in ()).throw(
            promptcard_runtime.ModelManagementError("document_input_not_supported")
        ),
    )
    monkeypatch.setattr(
        promptcard_runtime,
        "resolve_text_model",
        lambda binding: _descriptor(provider="deepseek", inputs=["text"]),
    )
    monkeypatch.setattr(promptcard_runtime, "_invoke_text_agent", invoke)
    monkeypatch.setattr(promptcard_runtime, "_load_canvas_image_attachments", _empty_attachments)
    body = promptcard_runtime.PromptCardRuntimeMessageRequest.model_validate({
        "content": "Read",
        "projectId": "project-1",
        "permissionScope": "prompt-library-agent",
        "documentResourceIds": ["pdf-1"],
    })

    with pytest.raises(HTTPException) as error:
        await promptcard_runtime.runtime_service.send_message(body, None)

    assert error.value.detail == "document_input_not_supported"
    assert invoked is False


@pytest.mark.anyio
async def test_text_documents_are_bounded_normal_user_input_for_non_ark_models(
    monkeypatch,
):
    from app.gateway import promptcard_runtime
    from app.gateway.ark_responses import ResolvedDocumentAsset

    async def load_documents(project_id, ids):
        return [ResolvedDocumentAsset(
            resource_id="text-1",
            filename="notes.txt",
            content_type="text/plain",
            content=b"normalized notes",
            normalized_text="normalized notes",
        )]

    captured = {}

    async def invoke(payload):
        captured.update(payload)
        return {"text": "ok", "canvasEdits": [], "proposals": []}

    descriptor = _descriptor(provider="deepseek", inputs=["text"])
    descriptor["model"]["integrationGroup"] = {
        "id": "pi-native",
        "displayName": "PI native",
        "kind": "pi-native",
    }
    monkeypatch.setattr(promptcard_runtime, "_load_document_resources", load_documents)
    monkeypatch.setattr(promptcard_runtime, "resolve_text_model", lambda binding: descriptor)
    monkeypatch.setattr(promptcard_runtime, "_invoke_text_agent", invoke)
    monkeypatch.setattr(promptcard_runtime, "_load_canvas_image_attachments", _empty_attachments)
    body = promptcard_runtime.PromptCardRuntimeMessageRequest.model_validate({
        "content": "Read",
        "projectId": "project-1",
        "permissionScope": "prompt-library-agent",
        "documentResourceIds": ["text-1"],
    })

    await promptcard_runtime.runtime_service.send_message(body, None)

    assert captured["content"] == (
        "Read\n\n[Attached document: notes.txt]\nnormalized notes"
    )
    assert "documentInvocationHandle" not in captured


@pytest.mark.anyio
async def test_no_document_internal_chat_keeps_original_completion_dispatch(monkeypatch):
    from app.gateway import promptcard_runtime

    calls = []

    def complete(payload, *, connection_id, model_id):
        calls.append((payload, connection_id, model_id))
        return {"content": [], "stopReason": "stop", "usage": {}}

    async def inline(function, *args, **kwargs):
        return function(*args, **kwargs)

    monkeypatch.setattr(promptcard_runtime, "complete_sdk_text", complete)
    monkeypatch.setattr(promptcard_runtime, "run_in_threadpool", inline)
    body = promptcard_runtime.PromptCardInternalChatRequest.model_validate({
        "connectionId": "connection-1",
        "model": "model-1",
        "systemPrompt": "system",
        "messages": [{"role": "user", "content": "hello"}],
        "tools": [],
    })

    result = await promptcard_runtime.runtime_service.internal_chat(body)

    assert result["stopReason"] == "stop"
    assert calls == [(
        body.model_dump(by_alias=True),
        "connection-1",
        "model-1",
    )]


@pytest.mark.anyio
async def test_internal_chat_resolves_opaque_handle_without_consuming_it(monkeypatch):
    from app.gateway import promptcard_runtime
    from app.gateway.ark_responses import ResolvedDocumentAsset

    assets = [
        ResolvedDocumentAsset(
            resource_id="text-1",
            filename="notes.txt",
            content_type="text/plain",
            content=b"normalized notes",
            normalized_text="normalized notes",
        ),
        ResolvedDocumentAsset(
            resource_id="pdf-1",
            filename="scan.pdf",
            content_type="application/pdf",
            content=b"%PDF-scan",
        ),
    ]
    handle = promptcard_runtime._document_invocations.register(assets, _descriptor())
    calls = []

    def complete(payload, *, connection_id, model_id, pdf_assets):
        calls.append((payload, connection_id, model_id, pdf_assets))
        return {"content": [], "stopReason": "stop", "usage": {}}

    async def inline(function, *args, **kwargs):
        return function(*args, **kwargs)

    monkeypatch.setattr(promptcard_runtime, "complete_sdk_text_with_documents", complete)
    monkeypatch.setattr(promptcard_runtime, "run_in_threadpool", inline)
    body = promptcard_runtime.PromptCardInternalChatRequest.model_validate({
        "connectionId": "connection-1",
        "model": "doubao-seed-2-0-lite-260215",
        "documentInvocationHandle": handle,
        "messages": [{"role": "user", "content": "hello"}],
    })

    await promptcard_runtime.runtime_service.internal_chat(body)
    await promptcard_runtime.runtime_service.internal_chat(body)

    assert len(calls) == 2
    payload = calls[0][0]
    assert "documentInvocationHandle" not in payload
    assert payload["messages"] == [{"role": "user", "content": "hello"}]
    assert calls[0][3] == [assets[1]]
    assert promptcard_runtime._document_invocations.resolve(handle) is not None
    promptcard_runtime._document_invocations.discard(handle)


@pytest.mark.anyio
async def test_internal_chat_rejects_forged_document_handle_before_completion(monkeypatch):
    from app.gateway import promptcard_runtime

    called = False

    def complete(*args, **kwargs):
        nonlocal called
        called = True

    monkeypatch.setattr(promptcard_runtime, "complete_sdk_text", complete)
    body = promptcard_runtime.PromptCardInternalChatRequest.model_validate({
        "connectionId": "connection-1",
        "model": "model-1",
        "documentInvocationHandle": "browser-forged-handle",
        "messages": [],
    })

    with pytest.raises(HTTPException) as error:
        await promptcard_runtime.runtime_service.internal_chat(body)

    assert error.value.detail == "document_context_unavailable"
    assert called is False


def test_startup_cleanup_failure_does_not_block_health_or_log_sensitive_values(
    monkeypatch,
    caplog,
):
    import importlib

    from fastapi.testclient import TestClient

    gateway_app = importlib.import_module("app.gateway.app")

    def fail_cleanup():
        raise RuntimeError(
            "opaque-remote-id C:/private/document.pdf secret-credential raw-body"
        )

    monkeypatch.setattr(gateway_app, "retry_provider_file_cleanup", fail_cleanup)

    with TestClient(gateway_app.create_app()) as client:
        response = client.get("/health")

    assert response.status_code == 200
    assert response.json()["status"] == "healthy"
    assert "opaque-remote-id" not in caplog.text
    assert "private/document.pdf" not in caplog.text
    assert "secret-credential" not in caplog.text
    assert "raw-body" not in caplog.text


@pytest.mark.anyio
async def test_lifespan_yields_health_before_a_stuck_cleanup_and_cancels_on_shutdown(
    monkeypatch,
):
    import asyncio
    import importlib

    gateway_app = importlib.import_module("app.gateway.app")
    started = threading.Event()
    release = threading.Event()
    worker_daemon: list[bool] = []

    def never_returns():
        worker_daemon.append(threading.current_thread().daemon)
        started.set()
        release.wait()

    monkeypatch.setattr(gateway_app, "retry_provider_file_cleanup", never_returns)
    context = gateway_app.lifespan(gateway_app.create_app())
    release_timer = threading.Timer(1.0, release.set)
    release_timer.start()
    try:
        await asyncio.wait_for(context.__aenter__(), timeout=0.2)
        assert await asyncio.to_thread(started.wait, 0.2)
        assert worker_daemon == [True]
        await asyncio.wait_for(context.__aexit__(None, None, None), timeout=0.2)
    finally:
        release.set()
        release_timer.cancel()


@pytest.mark.anyio
async def test_storage_document_seam_normalizes_txt_markdown_docx_and_keeps_pdf_bytes(
    monkeypatch,
):
    from io import BytesIO
    from urllib.parse import quote

    import httpx
    from docx import Document

    from app.gateway import promptcard_runtime

    document = Document()
    document.add_paragraph("DOCX paragraph")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "left"
    table.cell(0, 1).text = "right"
    docx_buffer = BytesIO()
    document.save(docx_buffer)
    fixtures = {
        "txt-1": ("notes.txt", "text/plain", b"line one\r\nline two"),
        "md-1": ("brief.md", "text/markdown", "Cafe\u0301".encode()),
        "docx-1": (
            "draft.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            docx_buffer.getvalue(),
        ),
        "pdf-1": ("scan.pdf", "application/pdf", b"%PDF-scanned-image-pages"),
    }
    requests: list[httpx.Request] = []

    def handler(request: httpx.Request):
        requests.append(request)
        resource_id = request.url.path.rsplit("/", 2)[-2]
        filename, content_type, content = fixtures[resource_id]
        return httpx.Response(
            200,
            content=content,
            headers={
                "Content-Type": content_type,
                "X-File-Name": quote(filename, safe=""),
                "X-Document-Resource-Id": resource_id,
            },
        )

    original_client = httpx.AsyncClient
    transport = httpx.MockTransport(handler)
    monkeypatch.setenv("PROMPTCARD_INTERNAL_TOKEN", "internal-token")
    monkeypatch.setattr(
        promptcard_runtime.httpx,
        "AsyncClient",
        lambda **kwargs: original_client(transport=transport, **kwargs),
    )

    assets = await promptcard_runtime._load_document_resources(
        "project-1",
        ["txt-1", "md-1", "docx-1", "pdf-1"],
    )

    assert [asset.normalized_text for asset in assets] == [
        "line one\nline two",
        "Caf\u00e9",
        "DOCX paragraph\nleft\tright",
        None,
    ]
    assert assets[-1].content == b"%PDF-scanned-image-pages"
    assert all(
        request.headers["X-PromptCard-Internal-Token"] == "internal-token"
        for request in requests
    )


@pytest.mark.anyio
async def test_cross_project_or_trashed_document_rejects_before_model_resolution(
    monkeypatch,
):
    import httpx

    from app.gateway import promptcard_runtime

    original_client = httpx.AsyncClient
    monkeypatch.setenv("PROMPTCARD_INTERNAL_TOKEN", "internal-token")
    monkeypatch.setattr(
        promptcard_runtime.httpx,
        "AsyncClient",
        lambda **kwargs: original_client(
            transport=httpx.MockTransport(
                lambda request: httpx.Response(404, json={"detail": "not_found"})
            ),
            **kwargs,
        ),
    )
    model_called = False

    def resolve_model(binding):
        nonlocal model_called
        model_called = True

    monkeypatch.setattr(promptcard_runtime, "resolve_text_model", resolve_model)
    body = promptcard_runtime.PromptCardRuntimeMessageRequest.model_validate({
        "content": "Read",
        "projectId": "project-1",
        "permissionScope": "prompt-library-agent",
        "documentResourceIds": ["other-project-or-trashed"],
    })

    with pytest.raises(HTTPException) as error:
        await promptcard_runtime.runtime_service.send_message(body, None)

    assert error.value.detail == "document_resource_invalid"
    assert model_called is False


@pytest.mark.anyio
async def test_storage_document_seam_rejects_path_like_filename_metadata(monkeypatch):
    import httpx

    from app.gateway import promptcard_runtime

    original_client = httpx.AsyncClient
    monkeypatch.setenv("PROMPTCARD_INTERNAL_TOKEN", "internal-token")
    monkeypatch.setattr(
        promptcard_runtime.httpx,
        "AsyncClient",
        lambda **kwargs: original_client(
            transport=httpx.MockTransport(
                lambda request: httpx.Response(
                    200,
                    content=b"safe text",
                    headers={
                        "Content-Type": "text/plain",
                        "X-File-Name": "C%3A%5Cprivate%5Cdocument.txt",
                        "X-Document-Resource-Id": "text-1",
                    },
                )
            ),
            **kwargs,
        ),
    )

    with pytest.raises(HTTPException) as error:
        await promptcard_runtime._load_document_resources(
            "project-1",
            ["text-1"],
        )

    assert error.value.detail == "document_storage_invalid_response"
